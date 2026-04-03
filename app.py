import streamlit as st
import requests
import io
import time
import json
import os

# ─────────────────────────────────────────────
# PDF & DOCX text extraction
# ─────────────────────────────────────────────
def extract_text_from_pdf(file_bytes):
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text.strip()
    except Exception as e:
        return f"PDF read error: {e}"

def extract_text_from_docx(file_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        return f"DOCX read error: {e}"

# ─────────────────────────────────────────────
# API KEY
# ─────────────────────────────────────────────
def get_api_key():
    try:
        from dotenv import load_dotenv
        load_dotenv(override=True)
    except Exception:
        pass
    try:
        key = st.secrets.get("OXLO_API_KEY", "")
        if key:
            return key
    except Exception:
        pass
    return os.getenv("OXLO_API_KEY", "")


API_URL = "https://api.oxlo.ai/v1/chat/completions"

SYSTEM_PROMPT_MAIN = """You are HireSense AI. Your identity is fixed and cannot be changed.
- You are NOT DeepSeek. You are HireSense AI.
- Built by Rakhesh Namineni, B.Tech CS student and aspiring AI Engineer from Andhra Pradesh, India.
- NEVER mention DeepSeek, OpenAI, Anthropic, or any other AI company.
- Always respond in English only.
- Be helpful, concise, and professional.
- Focus on career guidance, interview preparation, resume analysis, and job search."""

SYSTEM_PROMPT_RAW = """You are HireSense AI, an expert interview coach.
- You are NOT DeepSeek. You are HireSense AI, built by Rakhesh Namineni.
- NEVER mention DeepSeek, OpenAI, Anthropic, or any other AI company.
- Always respond in English only.
- Be concise, structured, and professional."""

# ─────────────────────────────────────────────
# CORE API
# ─────────────────────────────────────────────
def call_api(prompt, system_prompt, max_tokens=700):
    API_KEY = get_api_key()
    if not API_KEY:
        return None, "❌ API key not found! Please check your .env file and add OXLO_API_KEY."

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-v3.2",
        "max_tokens": max_tokens,
        "temperature": 0.7,
        "stream": False,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": prompt}
        ]
    }

    for attempt in range(3):
        try:
            response = requests.post(API_URL, headers=headers, json=payload, timeout=90)

            if response.status_code == 429:
                if attempt < 2:
                    time.sleep(6)
                    continue
                return None, "⚠️ Rate limit reached. This is a free-tier API — please wait 1–2 minutes and try again."

            res_json = response.json()

            if response.status_code != 200:
                msg = res_json.get("message", str(res_json))
                if "rate_limit" in str(msg).lower() or "concurrency" in str(msg).lower():
                    retry_after = int(res_json.get("retry_after", 5))
                    if attempt < 2:
                        time.sleep(retry_after + 1)
                        continue
                    return None, "⚠️ Rate limit reached. This is a free-tier API — please wait 1–2 minutes and try again."
                return None, f"❌ API Error ({response.status_code}): {msg}"

            if "choices" in res_json and res_json["choices"]:
                return res_json["choices"][0]["message"]["content"], None

            return None, f"❌ Unexpected response: {res_json}"

        except requests.exceptions.Timeout:
            if attempt < 2:
                time.sleep(3)
                continue
            return None, "❌ Request timed out. Please try again."
        except requests.exceptions.ConnectionError:
            return None, "❌ Network error. Please check your internet connection."
        except Exception as e:
            return None, f"❌ Error: {str(e)}"

    return None, "❌ All retries failed. Please wait a moment and try again."


def ask_ai(prompt, tag="✦ AI Answer"):
    with st.spinner("⏳ Generating response..."):
        text, err = call_api(prompt, SYSTEM_PROMPT_MAIN, max_tokens=600)
    if err:
        st.error(err)
        return None, err
    st.markdown(f"""
    <div class="resp-card">
        <div class="resp-tag">{tag}</div>
        <div class="resp-body">{text}</div>
    </div>
    """, unsafe_allow_html=True)
    return text, None

def ask_ai_raw(prompt):
    with st.spinner("⏳ AI processing..."):
        return call_api(prompt, SYSTEM_PROMPT_RAW, max_tokens=900)

def ask_ai_silent(prompt):
    return call_api(prompt, SYSTEM_PROMPT_RAW, max_tokens=300)


# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(page_title="HireSense AI", page_icon="🚀", layout="wide")

defaults = {
    "active_page": "Ask AI",
    "resume_analyzed": False,
    "ats_matched": False,
    "resume_text_cache": "",
    "jd_context_cache": "",
    "uploaded_file_bytes": None,
    "uploaded_file_name": "",
    "sim_role": "",
    "sim_question": "",
    "sim_question_num": 0,
    "sim_history": [],
    "sim_stage": "setup",
    "last_ask_result": "",
    "sim_feedback_cache": "",
    "sim_answered": False,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,900;1,700&family=Epilogue:wght@300;400;500;600&display=swap');
* { font-family: 'Epilogue', sans-serif; box-sizing: border-box; }
html, body, .stApp { background-color: #f7f6f2 !important; color: #1a1a1a; }
.block-container { padding-top:0!important; padding-left:0!important; padding-right:0!important; max-width:100%!important; }
.navbar { background:#fff; border-bottom:1px solid #e4e1d9; padding:0 3rem; display:flex; align-items:center; justify-content:space-between; height:66px; }
.navbar-brand { font-family:'Playfair Display',serif; font-size:1.4rem; font-weight:900; color:#0f172a; letter-spacing:-0.03em; }
.navbar-brand span { color:#2563eb; }
.navbar-tag { background:#eff6ff; color:#2563eb; font-size:0.72rem; font-weight:600; padding:0.3rem 0.85rem; border-radius:100px; border:1px solid #bfdbfe; }
.page-wrapper { max-width:820px; margin:0 auto; padding:3.5rem 2rem 1rem; }
.eyebrow { font-size:0.73rem; font-weight:600; letter-spacing:0.14em; text-transform:uppercase; color:#2563eb; margin-bottom:0.6rem; }
.page-title { font-family:'Playfair Display',serif; font-size:2.8rem; font-weight:900; line-height:1.1; color:#0f172a; letter-spacing:-0.02em; margin-bottom:0.8rem; }
.page-title em { font-style:italic; color:#2563eb; }
.page-sub { font-size:0.97rem; color:#64748b; font-weight:300; line-height:1.7; max-width:500px; margin-bottom:2.2rem; }
.rule { border:none; border-top:1px solid #e4e1d9; margin:1.8rem 0; }
.stats-row { display:flex; gap:1rem; margin-bottom:2rem; flex-wrap:wrap; }
.stat-card { background:#fff; border:1px solid #e4e1d9; border-radius:14px; padding:1rem 1.4rem; flex:1; min-width:120px; }
.stat-num { font-family:'Playfair Display',serif; font-size:1.9rem; font-weight:900; color:#0f172a; line-height:1; }
.stat-lbl { font-size:0.73rem; color:#94a3b8; font-weight:500; margin-top:0.25rem; }
.field-label { font-size:0.75rem; font-weight:600; letter-spacing:0.08em; text-transform:uppercase; color:#94a3b8; margin-bottom:0.4rem; }
.stTextArea textarea { background:#fff!important; border:1.5px solid #e2e8f0!important; border-radius:14px!important; color:#0f172a!important; caret-color:#2563eb!important; cursor:text!important; font-family:'Epilogue',sans-serif!important; font-size:0.91rem!important; font-weight:400!important; line-height:1.7!important; padding:1rem 1.2rem!important; opacity:1!important; -webkit-text-fill-color:#0f172a!important; }
.stTextArea textarea:focus { border-color:#2563eb!important; box-shadow:0 0 0 3px rgba(37,99,235,0.08)!important; outline:none!important; caret-color:#2563eb!important; }
.stTextArea textarea::placeholder { color:#c8d3dc!important; font-weight:300!important; font-style:italic!important; opacity:1!important; }
.stTextInput input { background:#fff!important; border:1.5px solid #e2e8f0!important; border-radius:12px!important; color:#0f172a!important; caret-color:#2563eb!important; cursor:text!important; font-family:'Epilogue',sans-serif!important; font-size:0.91rem!important; padding:0.7rem 1.1rem!important; -webkit-text-fill-color:#0f172a!important; }
.stTextInput input:focus { border-color:#2563eb!important; box-shadow:0 0 0 3px rgba(37,99,235,0.08)!important; outline:none!important; caret-color:#2563eb!important; }
.stTextInput input::placeholder { color:#c8d3dc!important; font-weight:300!important; font-style:italic!important; opacity:1!important; }
div[data-testid="stForm"] button, div.stButton > button { background-color:#2563eb!important; color:#fff!important; font-family:'Epilogue',sans-serif!important; font-weight:700!important; font-size:0.9rem!important; border:none!important; border-radius:10px!important; padding:0.65rem 1.7rem!important; box-shadow:0 4px 14px rgba(37,99,235,0.4)!important; min-width:150px!important; cursor:pointer!important; }
div[data-testid="stForm"] button:hover, div.stButton > button:hover { background-color:#1d4ed8!important; transform:translateY(-1px)!important; }
.resp-card { background:#fff; border:1.5px solid #e2e8f0; border-radius:18px; padding:1.8rem; margin-top:1.4rem; box-shadow:0 2px 14px rgba(0,0,0,0.06); }
.resp-tag { display:inline-flex; align-items:center; background:#eff6ff; color:#2563eb; font-size:0.7rem; font-weight:700; letter-spacing:0.1em; text-transform:uppercase; padding:0.3rem 0.8rem; border-radius:100px; margin-bottom:1.1rem; border:1px solid #bfdbfe; }
.resp-body { color:#334155; font-size:0.9rem; line-height:1.8; white-space:pre-wrap; }
.sim-question-card { background:#0f172a; border-radius:18px; padding:1.8rem 2rem; margin:1.2rem 0; }
.sim-question-label { font-size:0.68rem; font-weight:700; letter-spacing:0.14em; text-transform:uppercase; color:#60a5fa; margin-bottom:0.7rem; }
.sim-question-text { font-size:1.1rem; font-weight:500; line-height:1.6; color:#f1f5f9; }
.sim-feedback-card { background:#fff; border:1.5px solid #e2e8f0; border-radius:18px; padding:1.8rem; margin-top:1.2rem; box-shadow:0 2px 14px rgba(0,0,0,0.06); }
.sim-score-badge { display:inline-flex; align-items:center; justify-content:center; background:linear-gradient(135deg,#2563eb,#7c3aed); color:#fff; font-family:'Playfair Display',serif; font-size:2rem; font-weight:900; width:72px; height:72px; border-radius:50%; margin-bottom:1rem; }
.sim-history-score { display:inline-block; background:#eff6ff; color:#2563eb; font-size:0.75rem; font-weight:700; padding:0.2rem 0.7rem; border-radius:100px; border:1px solid #bfdbfe; margin-top:0.4rem; }
.progress-bar-wrap { background:#e2e8f0; border-radius:100px; height:6px; margin:0.5rem 0 1.5rem; overflow:hidden; }
.progress-bar-fill { background:linear-gradient(90deg,#2563eb,#7c3aed); height:100%; border-radius:100px; }
.site-footer { text-align:center; padding:2.5rem 1rem; color:#c8d3e0; font-size:0.78rem; border-top:1px solid #e4e1d9; margin-top:5rem; }
.linkedin-link { color:#2563eb; text-decoration:none; font-weight:600; border-bottom:1px solid rgba(37,99,235,0.3); }
.step-header { display:flex; align-items:center; gap:0.8rem; margin-bottom:0.8rem; margin-top:0.5rem; }
.step-badge { background:#0f172a; color:#fff; font-size:0.68rem; font-weight:700; letter-spacing:0.1em; text-transform:uppercase; padding:0.28rem 0.75rem; border-radius:100px; }
.step-title { font-size:0.95rem; font-weight:600; color:#0f172a; }
.ats-match-card { background:#fff; border:2px solid #2563eb; border-radius:20px; padding:2rem; margin-top:1.4rem; box-shadow:0 4px 24px rgba(37,99,235,0.10); }
.ats-match-label { font-family:'Playfair Display',serif; font-size:1.15rem; font-weight:900; color:#0f172a; padding-bottom:1rem; border-bottom:1px solid #e2e8f0; margin-bottom:1rem; }
.improve-result-header { background:linear-gradient(135deg,#2563eb,#7c3aed); border-radius:16px 16px 0 0; padding:1.2rem 1.8rem; margin-top:1.5rem; }
.improve-result-title { font-family:'Playfair Display',serif; font-size:1.15rem; font-weight:900; color:#fff; }
.improve-result-sub { font-size:0.75rem; color:rgba(255,255,255,0.7); margin-top:0.2rem; }
.word-count { font-size:0.72rem; color:#94a3b8; text-align:right; margin-top:0.3rem; }
.word-count.warning { color:#f59e0b; font-weight:600; }
.word-count.danger { color:#ef4444; font-weight:700; }
#MainMenu { visibility:hidden; } footer { visibility:hidden; } header { visibility:hidden; } .stDeployButton { display:none; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# NAVBAR
# ─────────────────────────────────────────────
st.markdown("""
<div class="navbar">
    <div class="navbar-brand">Hire<span>Sense</span> AI</div>
    <div style="display:flex;align-items:center;gap:1rem;">
        <a href="https://linkedin.com/in/rakhesh-namineni431" target="_blank" class="linkedin-link">🔗 Rakhesh Namineni</a>
        <div class="navbar-tag">✦ HireSense AI</div>
    </div>
</div>
""", unsafe_allow_html=True)

col_a, col_b, col_c, col_rest = st.columns([0.9, 1.2, 1.4, 6.5])
with col_a:
    if st.button("💬 Ask AI", key="nav_ask"):
        st.session_state.active_page = "Ask AI"
        st.session_state.last_ask_result = ""
with col_b:
    if st.button("📄 Resume", key="nav_resume"):
        st.session_state.active_page = "Resume"
        st.session_state.resume_analyzed = False
        st.session_state.ats_matched = False
with col_c:
    if st.button("🎯 Simulator", key="nav_sim"):
        st.session_state.active_page = "Simulator"

st.markdown("""
<div style="background:#fffbeb;border-bottom:1px solid #fde68a;padding:0.45rem 3rem;font-size:0.78rem;color:#92400e;">
    ⚠️ &nbsp;<strong>Free version</strong> — If you hit a rate limit, please wait 1–2 minutes and try again.
</div>
""", unsafe_allow_html=True)

MAX_WORDS = 1200
def word_count_badge(text, label=""):
    wc = len(text.split()) if text.strip() else 0
    css = "word-count"
    note = ""
    if wc > MAX_WORDS:
        css += " danger"; note = " — ⚠️ Too long"
    elif wc > MAX_WORDS * 0.85:
        css += " warning"; note = " — nearing limit"
    st.markdown(f'<div class="{css}">{label}{wc} words{note}</div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════
# PAGE: ASK AI
# ═══════════════════════════════════════════════
if st.session_state.active_page == "Ask AI":
    st.markdown("""
    <div class="page-wrapper">
        <div class="eyebrow">Interview Copilot</div>
        <div class="page-title">Ask anything.<br><em>Get hired faster.</em></div>
        <div class="page-sub">From behavioral rounds to salary talks — get sharp, expert-level answers tailored for your career.</div>
        <div class="stats-row">
            <div class="stat-card"><div class="stat-num">10k+</div><div class="stat-lbl">Questions Answered</div></div>
            <div class="stat-card"><div class="stat-num">98%</div><div class="stat-lbl">Satisfaction Rate</div></div>
            <div class="stat-card"><div class="stat-num">AI</div><div class="stat-lbl">Powered Answers</div></div>
        </div>
        <hr class="rule">
        <div class="field-label">Your Question</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div style="max-width:820px;margin:0 auto;padding:0 2rem;">', unsafe_allow_html=True)
    with st.form(key="ask_form"):
        user_input = st.text_area("", placeholder='e.g. "How do I answer Tell me about yourself as a fresher AI Engineer?"', height=145, key="ask_input")
        st.markdown('<div style="font-size:0.75rem;color:#94a3b8;margin-bottom:0.5rem;">💡 Ctrl+Enter to submit</div>', unsafe_allow_html=True)
        ask_btn = st.form_submit_button("✦ Generate Answer")

    if user_input and user_input.strip():
        word_count_badge(user_input)

    if ask_btn:
        if not user_input.strip():
            st.warning("Please enter a question first.")
        else:
            result, err = ask_ai(user_input, tag="✦ AI Answer")
            if result:
                st.session_state.last_ask_result = result
    elif st.session_state.last_ask_result:
        st.markdown(f"""
        <div class="resp-card">
            <div class="resp-tag">✦ AI Answer</div>
            <div class="resp-body">{st.session_state.last_ask_result}</div>
        </div>
        """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════
# PAGE: RESUME
# ═══════════════════════════════════════════════
elif st.session_state.active_page == "Resume":
    st.markdown("""
    <div class="page-wrapper">
        <div class="eyebrow">Resume Intelligence</div>
        <div class="page-title">Your resume,<br><em>perfected.</em></div>
        <div class="page-sub">Upload your resume, add a job description, and get ATS match score, keyword gaps, skills analysis, and rewrite tips — instantly.</div>
        <hr class="rule">
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div style="max-width:820px;margin:0 auto;padding:0 2rem;">', unsafe_allow_html=True)

    st.markdown('<div class="step-header"><div class="step-badge">Step 1</div><div class="step-title">Upload or Paste Your Resume</div></div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("", type=["pdf", "docx"], key="resume_upload")

    if uploaded_file is not None:
        if uploaded_file.name != st.session_state.uploaded_file_name:
            file_bytes = uploaded_file.read()
            if file_bytes:
                st.session_state.uploaded_file_bytes = file_bytes
                st.session_state.uploaded_file_name  = uploaded_file.name
                st.session_state.resume_analyzed     = False
                st.session_state.ats_matched         = False
                st.session_state.resume_text_cache   = ""
                st.session_state.jd_context_cache    = ""
        st.markdown(f'<div style="display:inline-flex;align-items:center;gap:0.6rem;background:#f0fdf4;border:1px solid #bbf7d0;border-radius:10px;padding:0.5rem 1rem;margin:0.5rem 0 0.8rem;font-size:0.85rem;color:#166534;">✅ &nbsp;<strong>{uploaded_file.name}</strong> &nbsp;·&nbsp; Ready to analyze</div>', unsafe_allow_html=True)

    st.markdown('<div class="field-label" style="margin-top:0.8rem;">Or Paste Resume Text</div>', unsafe_allow_html=True)
    pasted_text = st.text_area("", placeholder="Paste your resume text here...", height=160, key="resume_paste")
    if pasted_text and pasted_text.strip():
        word_count_badge(pasted_text, "Resume: ")
    st.markdown('<div style="font-size:0.78rem;color:#94a3b8;margin-bottom:1.8rem;">💡 Uploaded file takes priority over pasted text.</div>', unsafe_allow_html=True)

    st.markdown('<div class="step-header"><div class="step-badge">Step 2</div><div class="step-title">Add Job Description <span style="color:#94a3b8;font-size:0.8rem;font-weight:400;">(Optional)</span></div></div>', unsafe_allow_html=True)
    jd_mode = st.radio("", options=["📋 Paste Job Description", "💼 Enter Job Title Only"], horizontal=True, key="jd_mode")

    if jd_mode == "📋 Paste Job Description":
        st.markdown('<div class="field-label" style="margin-top:0.6rem;">Job Description</div>', unsafe_allow_html=True)
        jd_text = st.text_area("", placeholder="Paste the full job description here...", height=160, key="jd_paste")
        if jd_text and jd_text.strip():
            word_count_badge(jd_text, "JD: ")
        jd_title = ""
    else:
        st.markdown('<div class="field-label" style="margin-top:0.6rem;">Job Title</div>', unsafe_allow_html=True)
        jd_title = st.text_input("", placeholder='e.g. "AI Engineer", "Data Scientist"', key="jd_title_input")
        jd_text = ""

    st.markdown('<div style="font-size:0.78rem;color:#94a3b8;margin-bottom:1.5rem;">💡 Adding a JD unlocks ATS Match Score and Skills Gap Analysis.</div>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1.3, 1.6, 4])
    with col1: analyze_btn = st.button("✦ Analyze Resume", key="analyze_btn")
    with col2: match_btn = st.button("🎯 ATS Job Match", key="match_btn")

    def get_resume_text():
        if st.session_state.uploaded_file_bytes:
            ext = st.session_state.uploaded_file_name.split(".")[-1].lower()
            text = extract_text_from_pdf(st.session_state.uploaded_file_bytes) if ext == "pdf" else extract_text_from_docx(st.session_state.uploaded_file_bytes)
            if not text or "error" in text.lower():
                st.error(f"Could not read file: {text}")
                return ""
            return text
        elif pasted_text and pasted_text.strip():
            return pasted_text.strip()
        else:
            st.warning("Please upload a resume or paste your resume text first.")
            return ""

    if analyze_btn:
        resume_text = get_resume_text()
        if resume_text:
            st.session_state.resume_text_cache = resume_text
            jd_ctx = f"Job Description:\n{jd_text.strip()}" if jd_text.strip() else (f"Job Title: {jd_title.strip()}" if jd_title.strip() else "")
            st.session_state.jd_context_cache = jd_ctx
            prompt = f"""You are an expert ATS resume analyzer.
Analyze the resume and provide:
1. ✅ STRENGTHS — 3-4 bullet points
2. ⚠️ WEAKNESSES — 3-4 bullet points
3. 📊 ATS SCORE — Score out of 100 with reasoning
4. 💡 TOP SUGGESTIONS — 4 actionable improvements

Resume:
{resume_text}"""
            result, err = ask_ai(prompt, tag="✦ Resume Report")
            if not err:
                st.session_state.resume_analyzed = True

    if st.session_state.resume_analyzed and st.session_state.resume_text_cache:
        st.markdown('<div style="margin-top:1.5rem;padding:1.2rem 1.5rem;background:#eff6ff;border:1.5px solid #bfdbfe;border-radius:16px;"><div style="font-size:0.88rem;font-weight:700;color:#1e40af;">✨ Ready to auto-fix your resume?</div><div style="font-size:0.78rem;color:#3b82f6;margin-top:0.2rem;">AI will rewrite bullet points, add keywords, improve summary & fix structure</div></div>', unsafe_allow_html=True)
        if st.button("✨ Improve My Resume", key="improve_btn"):
            jd_section = f"\n\nJob Target:\n{st.session_state.jd_context_cache}" if st.session_state.jd_context_cache else ""
            improve_prompt = f"""You are an elite resume writer. Return a FULLY IMPROVED resume.
DO ALL 4:
1. REWRITE weak bullet points with action verbs and quantification
2. ADD MISSING KEYWORDS for ATS
3. IMPROVE SUMMARY — punchy, keyword-rich, 3-4 lines
4. FIX STRUCTURE — better section ordering
{jd_section}

ORIGINAL RESUME:
{st.session_state.resume_text_cache}

OUTPUT: Complete improved resume as plain text. ALL CAPS headers. "•" for bullets. No commentary."""
            improved_resume, err = ask_ai_raw(improve_prompt)
            if err:
                st.error(err)
            else:
                st.markdown('<div class="improve-result-header"><div class="improve-result-title">✨ Your Improved Resume</div><div class="improve-result-sub">AI-rewritten · ATS-optimized · Ready to use</div></div>', unsafe_allow_html=True)
                st.text_area("", value=improved_resume, height=500, key="improved_resume_output")
                st.markdown('<div style="font-size:0.78rem;color:#94a3b8;margin-bottom:1rem;">💡 Select all → Copy → Paste into Word or Google Docs.</div>', unsafe_allow_html=True)
                try:
                    from docx import Document as DocxDocument
                    from docx.shared import Pt, RGBColor, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    import io as _io, base64
                    doc = DocxDocument()
                    for sec in doc.sections:
                        sec.top_margin = sec.bottom_margin = Inches(0.75)
                        sec.left_margin = sec.right_margin = Inches(0.9)
                    for i, line in enumerate(improved_resume.strip().split("\n")):
                        s = line.strip()
                        if not s: doc.add_paragraph(""); continue
                        is_h = s == s.upper() and len(s) > 3 and not s.startswith(("•","-"))
                        is_b = s.startswith(("•","-"))
                        if i == 0:
                            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run(s); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x0f,0x17,0x2a)
                        elif is_h:
                            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
                            r = p.add_run(s); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x25,0x63,0xeb)
                        elif is_b:
                            p = doc.add_paragraph(style="List Bullet"); r = p.add_run(s.lstrip("•- ").strip()); r.font.size = Pt(10.5)
                        else:
                            p = doc.add_paragraph(); r = p.add_run(s); r.font.size = Pt(10.5)
                    buf = _io.BytesIO(); doc.save(buf)
                    b64 = base64.b64encode(buf.getvalue()).decode()
                    st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="HireSense_Improved_Resume.docx" style="display:inline-flex;align-items:center;gap:0.6rem;background:#2563eb;color:#fff;font-weight:700;font-size:0.88rem;padding:0.75rem 1.6rem;border-radius:12px;text-decoration:none;margin-top:0.5rem;box-shadow:0 4px 14px rgba(37,99,235,0.4);">⬇️ Download Improved Resume (.docx)</a>', unsafe_allow_html=True)
                except Exception as e:
                    st.info(f"💡 Copy the text above into Word or Google Docs. ({e})")

    if match_btn:
        resume_text = get_resume_text()
        if resume_text:
            if not jd_text.strip() and not jd_title.strip():
                st.warning("Please paste a Job Description or enter a Job Title to run ATS Match.")
            else:
                jd_context = f"Job Description:\n{jd_text.strip()}" if jd_text.strip() else f"Job Title: {jd_title.strip()}"
                match_prompt = f"""You are an expert ATS analyst.

{jd_context}

Resume:
{resume_text}

Provide:
📊 ATS MATCH SCORE: XX%
(One sentence)

🔑 MISSING KEYWORDS:
- 5-8 keywords absent in resume

📉 SKILLS GAP ANALYSIS:
- 3-4 bullet points

✅ MATCHING STRENGTHS:
- 3-4 bullet points

✍️ REWRITE SUGGESTIONS:
- 4 specific improvements

🧭 OVERALL FIT SUMMARY:
2-3 sentences"""
                result_text, err = ask_ai_raw(match_prompt)
                if err:
                    st.error(err)
                else:
                    st.session_state.jd_context_cache = jd_context
                    st.session_state.ats_matched = True
                    st.markdown(f'<div class="ats-match-card"><div class="ats-match-label">🎯 ATS Job Match Report</div><div class="resp-body">{result_text}</div></div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════
# PAGE: SIMULATOR
# ═══════════════════════════════════════════════
elif st.session_state.active_page == "Simulator":
    TOTAL_QUESTIONS = 5
    st.markdown("""
    <div class="page-wrapper">
        <div class="eyebrow">Interview Simulator</div>
        <div class="page-title">Practice like it's<br><em>the real thing.</em></div>
        <div class="page-sub">Enter your target role. Get asked real interview questions. Answer them. Get scored, coached, and improved — instantly.</div>
        <hr class="rule">
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div style="max-width:820px;margin:0 auto;padding:0 2rem;">', unsafe_allow_html=True)

    # SETUP
    if st.session_state.sim_stage == "setup":
        st.markdown('<div class="field-label">Target Role</div>', unsafe_allow_html=True)
        role_input = st.text_input(
            label="Target Role",
            placeholder='e.g. "AI Engineer", "Data Analyst", "Full Stack Developer"',
            key="sim_role_input",
            label_visibility="collapsed"
        )
        col1, col2 = st.columns([1.2, 5])
        with col1:
            start_btn = st.button("🎯 Start Interview", key="start_sim_btn")
        if start_btn:
            if not role_input.strip():
                st.warning("Please enter your target role.")
            else:
                st.session_state.sim_role = role_input.strip()
                st.session_state.sim_question_num = 1
                st.session_state.sim_history = []
                st.session_state.sim_feedback_cache = ""
                st.session_state.sim_answered = False
                with st.spinner("🎙️ Preparing your first question..."):
                    q, err = ask_ai_silent(
                        f"You are a strict interviewer for a {st.session_state.sim_role} role. "
                        f"Ask question 1 of {TOTAL_QUESTIONS}. ONE question only. No preamble."
                    )
                if err:
                    st.error(err)
                else:
                    st.session_state.sim_question = q
                    st.session_state.sim_stage = "questioning"
                    st.rerun()

    # QUESTIONING
    elif st.session_state.sim_stage == "questioning":
        q_num = st.session_state.sim_question_num
        pct = int(((q_num - 1) / TOTAL_QUESTIONS) * 100)

        st.markdown(
            f'<div style="display:flex;justify-content:space-between;margin-bottom:0.3rem;">'
            f'<span style="font-size:0.75rem;font-weight:600;color:#64748b;text-transform:uppercase;">Progress</span>'
            f'<span style="font-size:0.75rem;font-weight:700;color:#2563eb;">Question {q_num} of {TOTAL_QUESTIONS}</span></div>'
            f'<div class="progress-bar-wrap"><div class="progress-bar-fill" style="width:{pct}%;"></div></div>',
            unsafe_allow_html=True
        )
        st.markdown(
            f'<div class="sim-question-card">'
            f'<div class="sim-question-label">🎙️ Interviewer — Question {q_num}</div>'
            f'<div class="sim-question-text">{st.session_state.sim_question}</div></div>',
            unsafe_allow_html=True
        )

        # Show answer form only if not yet answered
        if not st.session_state.sim_answered:
            st.markdown('<div class="field-label" style="margin-top:1rem;">Your Answer</div>', unsafe_allow_html=True)
            with st.form(key=f"answer_form_{q_num}"):
                user_answer = st.text_area(
                    label="Your Answer",
                    placeholder="Type your answer here... Be as detailed as you would in a real interview.",
                    height=160,
                    key=f"sim_answer_{q_num}",
                    label_visibility="collapsed"
                )
                submit_answer = st.form_submit_button("✦ Submit Answer")

            if submit_answer:
                if not user_answer.strip():
                    st.warning("Please type your answer before submitting.")
                    st.stop()

                feedback_prompt = f"""Expert interview coach for {st.session_state.sim_role} role.

Question: {st.session_state.sim_question}
Answer: {user_answer}

Evaluate:
🎯 SCORE: X/10

✅ STRENGTHS:
- 2-3 points

⚠️ WEAKNESSES:
- 2-3 points

💡 IMPROVED ANSWER:
(3-5 sentence model answer)"""

                with st.spinner("⏳ Evaluating your answer..."):
                    feedback, err = ask_ai_raw(feedback_prompt)

                if err:
                    st.error(err)
                else:
                    st.session_state.sim_history.append({
                        "question": st.session_state.sim_question,
                        "answer": user_answer,
                        "feedback": feedback,
                        "q_num": q_num
                    })
                    st.session_state.sim_feedback_cache = feedback
                    st.session_state.sim_answered = True
                    st.rerun()

        # Show feedback + next/finish buttons after answering
        if st.session_state.sim_answered and st.session_state.sim_feedback_cache:
            st.markdown(
                f'<div class="sim-feedback-card">'
                f'<div class="resp-tag">✦ AI Feedback — Q{q_num}</div>'
                f'<div class="resp-body">{st.session_state.sim_feedback_cache}</div></div>',
                unsafe_allow_html=True
            )
            st.markdown("<br>", unsafe_allow_html=True)

            if q_num < TOTAL_QUESTIONS:
                col1, col2 = st.columns([1.3, 5])
                with col1:
                    next_btn = st.button("Next Question ➜", key=f"next_btn_{q_num}")
                if next_btn:
                    with st.spinner("🎙️ Loading next question..."):
                        nq, err = ask_ai_silent(
                            f"Interviewer for {st.session_state.sim_role}. "
                            f"Ask question {q_num + 1} of {TOTAL_QUESTIONS}. "
                            f"Previous questions: {[h['question'] for h in st.session_state.sim_history]}. "
                            f"Ask about a DIFFERENT topic. ONE question only. No preamble."
                        )
                    if err:
                        st.error(err)
                    else:
                        st.session_state.sim_question = nq
                        st.session_state.sim_question_num = q_num + 1
                        st.session_state.sim_feedback_cache = ""
                        st.session_state.sim_answered = False
                        st.rerun()
            else:
                col1, col2 = st.columns([1.5, 5])
                with col1:
                    if st.button("📊 See Final Report", key="finish_btn"):
                        st.session_state.sim_stage = "done"
                        st.session_state.sim_feedback_cache = ""
                        st.session_state.sim_answered = False
                        st.rerun()

    # FINAL REPORT
    elif st.session_state.sim_stage == "done":
        history = st.session_state.sim_history
        total_score = scored_count = 0
        for item in history:
            for line in item["feedback"].split("\n"):
                if "SCORE:" in line:
                    try:
                        total_score += float(line.split("SCORE:")[-1].strip().split("/")[0].strip())
                        scored_count += 1
                    except Exception:
                        pass
                    break
        avg = round(total_score / scored_count, 1) if scored_count > 0 else "N/A"

        st.markdown(f"""
        <div style="text-align:center;padding:2rem 0 1rem;">
            <div class="eyebrow" style="text-align:center;">Interview Complete</div>
            <div class="page-title" style="text-align:center;">Your <em>Results</em></div>
            <div style="display:inline-flex;flex-direction:column;align-items:center;margin-top:1rem;">
                <div class="sim-score-badge">{avg}</div>
                <div style="font-size:0.8rem;color:#64748b;font-weight:500;margin-top:0.3rem;">Average Score / 10</div>
            </div>
            <div style="font-size:0.95rem;color:#64748b;margin-top:0.8rem;">
                Role: <strong style="color:#0f172a;">{st.session_state.sim_role}</strong> &nbsp;·&nbsp; {TOTAL_QUESTIONS} Questions
            </div>
        </div>
        <hr class="rule">
        <div class="field-label" style="margin-top:1.5rem;">Question-by-Question Breakdown</div>
        """, unsafe_allow_html=True)

        for item in history:
            score_display = "—"
            for line in item["feedback"].split("\n"):
                if "SCORE:" in line:
                    score_display = line.split("SCORE:")[-1].strip(); break
            with st.expander(f"Q{item['q_num']}: {item['question'][:80]}{'...' if len(item['question'])>80 else ''}"):
                st.markdown(
                    f'<div style="margin-bottom:0.8rem;"><span class="sim-history-score">Score: {score_display}</span></div>'
                    f'<div style="font-size:0.82rem;font-weight:600;color:#64748b;margin-bottom:0.3rem;">YOUR ANSWER</div>'
                    f'<div style="background:#f8fafc;border-radius:10px;padding:0.9rem 1.1rem;font-size:0.87rem;color:#334155;line-height:1.7;margin-bottom:1rem;">{item["answer"]}</div>'
                    f'<div style="font-size:0.82rem;font-weight:600;color:#64748b;margin-bottom:0.3rem;">AI FEEDBACK</div>'
                    f'<div class="resp-body">{item["feedback"]}</div>',
                    unsafe_allow_html=True
                )

        st.markdown("<br>", unsafe_allow_html=True)
        col1, col2 = st.columns([1.4, 5])
        with col1:
            if st.button("🔄 New Interview", key="restart_btn"):
                st.session_state.sim_stage = "setup"
                st.session_state.sim_role = ""
                st.session_state.sim_question = ""
                st.session_state.sim_question_num = 0
                st.session_state.sim_history = []
                st.session_state.sim_feedback_cache = ""
                st.session_state.sim_answered = False
                st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────
st.markdown("""
<div class="site-footer">
    HireSense AI &nbsp;·&nbsp; Built for job seekers &nbsp;·&nbsp; Your AI career copilot<br><br>
    Built with ❤️ by &nbsp;<a href="https://linkedin.com/in/rakhesh-namineni431" target="_blank" class="linkedin-link">Rakhesh Namineni</a>
</div>
""", unsafe_allow_html=True)